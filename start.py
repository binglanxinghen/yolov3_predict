import datetime
from flask import Flask
from flask import render_template, redirect, url_for, request, session
from werkzeug import secure_filename
from flask import send_file, send_from_directory
from flask import make_response
import os
import shutil
from flask_paginate import Pagination,get_page_parameter
import time
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash  # 密码保护，使用hash方法
from flask_sqlalchemy import SQLAlchemy

from docx import Document
from docx.shared import Inches
from PIL import Image
from predict import YoloTest
# import python4
basedir = os.path.abspath(os.path.dirname(__file__)) #定义一个根目录 用于保存图片用

os.environ["CUDA_VISIBLE_DEVICES"] = '1'


yolo_mini = YoloTest()

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///students.sqlite3'
app.config['SECRET_KEY'] = "random string"

db = SQLAlchemy(app)

#信息表
class Info(db.Model):
    __tablename__ = 'info'
    info_id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    username = db.Column(db.String(20), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    signature= db.Column(db.String(255),default='')
    telphone = db.Column(db.String(225),default='')
    like_flower=db.Column(db.String(225))
    sex=db.Column(db.String(10),default='')
    age=db.Column(db.Integer,default='')
    img_name=db.Column(db.String(10),nullable=False)
    create_time = db.Column(db.DateTime, default=datetime.now())

#user表
class User(db.Model):
    __tablename__ = 'user'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    email = db.Column(db.String(225), nullable=False)
    _password = db.Column(db.String(200), nullable=False)  # 内部使用

    @property
    def password(self):  # 定义一个外部使用的密码
        return self._password

    @password.setter  # 设置密码加密
    def password(self, row_password):
        self._password = generate_password_hash(row_password)

    # 检查密码是否正确
    def check_password(self, row_password):  # 定义一个反向解密的函数
        result = check_password_hash(self._password, row_password)
        return result

#db.drop_all()
#print('delete success')
#db.create_all()
#print('create success')

class History(db.Model):
    __tablename__ = 'history'
    history_id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    result = db.Column(db.Text, nullable=False)
    file_name = db.Column(db.String(225),nullable=False)
    time = db.Column(db.DateTime, default=datetime.now())

def count_classes(flow_list):
    classes = {"daisy":0,"dandelion":0,"rose":0,"sunflowers":0,"tulips":0}
    sum=0
    for item in flow_list:
        if item.split()[0] in classes.keys():
            classes[item.split()[0]]= classes[item.split()[0]] + 1
            sum=sum+1
    if(sum==0):
        s="Unknown"
        return s
    s=''
    for key,value in classes.items():
        if(value!=0):
            if(key=='daisy'):
                key='雏菊'
            if(key=='dandelion'):
                key='蒲公英'
            if(key=='rose'):
                key='玫瑰花'
            if(key=='sunflowers'):
                key='向日葵'
            if(key=='tulips'):
                key='郁金香'
            s=s+(key+":"+str(value)+"\n")
    return s

@app.route("/download/<history_id>", methods=["GET"])
def download(history_id):
    history = History.query.filter(History.history_id == history_id).first()
    images = "static/result/"+str(history.user_id)+"/"+history.file_name  # 保存在本地的图片
    doc = Document()
    doc.add_heading('结果报告', level=2)
    try:
        doc.add_picture(images, width=Inches(2))  # 添加图, 设置宽度
    except Exception:
        jpg_ima = Image.open(images)  # 打开图片
        jpg_ima.save('0.jpg')  # 保存新的图片
        doc.add_picture(images, width=Inches(2))  # 添加图, 设置宽度
    doc.add_paragraph("结果：")  # 添加文字
    doc.add_paragraph(history.result)  # 添加文字
    doc.add_paragraph("时间：" + str(history.time))  # 添加文字
    doc.add_paragraph("本报告来自智能识别花朵网站")  # 添加文字
    doc.save('static/result.docx')  # 保存路径
    directory = os.getcwd()  # 假设在当前目录
    response = make_response(send_from_directory(directory+"/static","result.docx" , as_attachment=True))
    response.headers["Content-Disposition"] = "attachment; filename={}".format("result.docx" .encode().decode('latin-1'))
    return response

#调用predict
# @app.route("/forecast", methods=["POST"])

@app.route("/forecast", methods=["POST"])
def forecast(path,filename):
    a=time.time()
    result_str=""
    if not session.get('id'):
        #游客
        yolo_mini.evaluate("./" + path + filename, "./" + path + filename)
        # results = os.system(
        #     "python predict.py %s %s %s %s" % ("-i", "./" + path + filename, "-o", "./" + path + filename))
    else:
        #登录用户
        yolo_mini.evaluate( "./" + path + filename, "./static/result/" + str(session.get('id')) + "/" + filename)
        # results = os.system(
        #     "python predict.py %s %s %s %s" % ("-i", "./" + path + filename, "-o", "./static/result/" + str(session.get('id')) + "/" + filename))
    with open("flow_result.txt",'r') as f:
        lines = f.readlines()
        result_str = count_classes(lines)
    b=time.time()
    print("All time:%f"%(b-a))
    return result_str

#主页
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/history')
def history():
    if not session.get('id'):
        return render_template('login.html')
    id=session.get('id')
    page = request.args.get('page', 1, type=int)
    pagination = History.query.filter(History.user_id == id).order_by(History.time.desc()).paginate(page, per_page=3, error_out=False)
    historys = pagination.items
    return render_template('history.html', historys=historys, pagination=pagination,base_url="img/"+str(id)+"/")

def add_history(user_id,fname,result):
    history = History(result=result, file_name=fname, user_id=user_id)
    db.session.add(history)
    db.session.flush()
    hid = history.history_id
    db.session.commit()
    db.session.flush()
    return hid

#删除历史纪录
@app.route('/delete/<history_id>')
def del_history(history_id):
    if not session.get('id'):
        return redirect(url_for('login'))
    else:
        history = History.query.filter(History.history_id == history_id).first()
        if session.get('id')!=history.user_id:
            return redirect(url_for('login'))
        else:
            db.session.delete(history)
            db.session.commit()
            db.session.flush()
        return redirect(url_for('history'))

#拍照上传
@app.route('/camera',methods=['GET','POST'])
def camera():
    print('camera')
    if request.method == 'POST':
        print('camera')
        f = request.files['file1']
        basepath = os.path.dirname(__file__)  # 当前文件所在路径
        fname= time.strftime("%Y%m%d%H%M%S",time.localtime(time.time()))+secure_filename(f.filename)
        if not session.get('id'):
            # 游客
            print('游客')
            upload_path = os.path.join(basepath, 'static/image/', fname)
            # 游客的放在static/image/
            f.save(upload_path)
            print(upload_path)
            res = forecast('static/image/', fname)
            return render_template('detail.html',image_url="image/"+fname,result=res,time=datetime.now(),history=None)
        dirs = 'img/' + str(session.get('id'))+'/'
        path = 'static/'+dirs
        # print(path)
        if not os.path.exists(path):
            print(path)
            os.makedirs(path)
        upload_path = os.path.join(basepath, path, fname)
        # 保存文件位置static/img/<user_id>/
        f.save(upload_path)
        #获取结果
        result = forecast(path,fname)
        hid = add_history(session.get('id'), fname, result)
        # print(hid)
        return redirect(url_for('detail', history_id=hid))
    return render_template('index.html')

#上传文件
@app.route('/upload',methods=['GET','POST'])
def upload():
    print('upload')
    if request.method == 'POST':
        print('upload')
        f=request.files['file2']
        basepath = os.path.dirname(__file__)  # 当前文件所在路径
        fname = time.strftime("%Y%m%d%H%M%S", time.localtime(time.time())) + secure_filename(f.filename)
        if not session.get('id'):
            # 游客
            print('游客')
            upload_path = os.path.join(basepath, 'static/image/', fname)
            # 游客的放在static/image/
            f.save(upload_path)
            print(upload_path)
            res = forecast('static/image/', fname)
            return render_template('detail.html', image_url="image/"+fname, result=res, time=datetime.now(), history=None)
        dirs = 'img/' + str(session.get('id')) + '/'
        path = 'static/' + dirs
        # print(path)
        if not os.path.exists(path):
            print(path)
            os.makedirs(path)
        upload_path = os.path.join(basepath, path, fname)
        # 保存文件位置static/img/<user_id>/
        f.save(upload_path)
        # 获取结果
        result = forecast(path, fname)
        hid = add_history(session.get('id'), fname, result)
        # print(hid)
        return redirect(url_for('detail', history_id=hid))
    return render_template('index.html')

#时间格式化
@app.template_filter('datetime_format')
def _jinja2_filter_datetime_format(datetimeValue, format='%y/%m/%d %H:%M:%S'):
    """convert a datetime to a different format."""
    print(datetime)
    return datetimeValue.strftime(format)

#详情
@app.route('/detail/<history_id>')
def detail(history_id):
    if not session.get('id'):
        return render_template('login.html')
    history = History.query.filter(History.history_id == history_id).first()
    return render_template('detail.html',image_url="result/"+str(history.user_id)+"/"+history.file_name, history = history)

#关于
@app.route('/about')
def about():
    return render_template('about.html')

#登录
@app.route('/login/',methods=['GET','POST'])
def login():
    if request.method=='GET':
        return render_template('login.html')
    else:
        email = request.form.get('email')
        password = request.form.get('password')
        print(email)
        print(password)
        user = User.query.filter(User.email == email).first()
        print(user)
        if user:
            if user.check_password(password):
                user_info = Info.query.filter(Info.user_id == user.id).first()
                session['user']=user_info.username
                session['id']=user.id
                session['head_img_url']='img/'+str(user.id)+'/'+user_info.img_name
                session.permanent = True
                return redirect(url_for('index'))
            else:
                return render_template('login.html', error='user password error')
        else:
            return render_template('login.html',error="username not exited")

#退出
@app.route('/logout/',methods=['GET','POST'])
def logout():
    session.clear()
    return redirect(url_for('index'))

#百科
@app.route('/more/',methods=['GET','POST'])
def more():
    return render_template('more.html')

#修改头像
@app.route('/editImg/<user_id>/',methods=['POST'])
def editImg(user_id):
    print("更改头像")
    f = request.files['file']
    basepath = os.path.dirname(__file__)  # 当前文件所在路径
    fname = 'head'+time.strftime("%Y%m%d%H%M%S", time.localtime(time.time()))+'.jpg'
    upload_path = os.path.join(basepath, 'static/img/'+str(user_id), fname)
    f.save(upload_path)
    userInfo = Info.query.filter(Info.user_id == user_id).first()
    userInfo.img_name = fname
    db.session.commit()
    return redirect(url_for('info',user_id=user_id))

#修改资料
@app.route('/editInfo/<user_id>/',methods=['POST'])
def editInfo(user_id):
    username = request.form.get('name')
    print(username)
    sex = request.form.get('sex')
    signature = request.form.get('signature')
    telphone = request.form.get('telphone')
    age = request.form.get('age')
    userInfo = Info.query.filter(Info.user_id == user_id).first()
    userInfo.username=username
    userInfo.sex=sex
    userInfo.age=age
    userInfo.signature=signature
    userInfo.telphone=telphone
    db.session.commit()
    session['user']=username
    print(session['user'])
    return redirect(url_for('info',user_id=user_id))

#退出
@app.route('/info/<user_id>/',methods=['GET','POST'])
def info(user_id):
    userInfo = Info.query.filter(Info.user_id == user_id).first()
    user=User.query.filter(User.id==user_id).first()
    print(userInfo)
    return render_template('info.html',userInfo=userInfo,head_img_url='img/'+str(session.get('id'))+'/'+userInfo.img_name,user_email=user.email)

def copyImg(source,path,filename):
    shutil.copyfile(source,path+filename)

#注册
@app.route('/register',methods=['GET','POST'])
def register():
    if request.method=='GET':
        return render_template('register.html')
    else:
        username = request.form.get('username')
        password = request.form.get('password')
        email = request.form.get('email')
        user = User.query.filter(User.email==email).first()
        if user:
            return render_template('register.html',error='该邮箱已被注册')
        else:
            user = User(email=email,password=password)
            db.session.add(user)
            db.session.flush()
            user_id = user.id
            db.session.commit()
            user_info = Info(user_id=user_id,username=username,img_name='default_head1.jpg')
            print(email)
            db.session.add(user_info)
            db.session.flush()
            print(user_info)
            db.session.commit()
            #创建文件夹
            path1 = 'static/result/'+str(user_id)
            path2 = 'static/img/' + str(user_id)
            os.makedirs(path1)
            os.makedirs(path2)
            copyImg('static/picture/default_head1.jpg', path2+'/','default_head1.jpg')
            return redirect(url_for('login'))

# 修改密码
@app.route('/edit_password/', methods=['GET', 'POST'])
def edit_password():
    if request.method == 'GET':
        return render_template("edit_password.html")
    else:
        newpassword = request.form.get('password')
        user = User.query.filter(User.username == request.form.get('username')).first()
        if user:
            user.password = newpassword
            db.session.commit()
        else:
            return 'not existed'
        return redirect(url_for('index'))


if __name__ =="__main__":
     app.run(host='0.0.0.0',port=8080)
